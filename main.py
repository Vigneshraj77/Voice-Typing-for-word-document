from tkinter import *
import speech_recognition as sr 
import pyttsx3
from docx import Document
from docx.shared import Inches
from tkinter import filedialog

#function to browse the loction of address
def browse_button():
    filename = filedialog.askdirectory()
    print(filename)
    global i
    i.set(filename)
    return

doc = Document()
doc_para = doc.add_paragraph("")

#function to insert image into document
def insert():
    global doc
    global i
    address=i.get()
    doc.add_picture(address,width=Inches(1.5))
    return

#saving the word document
def save():
    global doc
    global name
    temp= "C:/Users/Hp/Documents/"
    string=name.get()
    temp=temp+string+".docx"
    print(temp)
    doc.save(temp)
    return

#creating the interface
myWindow=Tk()
myWindow.title("Voice Typing")
myWindow.geometry("700x500")
#function for listening to voice
def listen():
    def write(txt):
        global doc
        doc_para = doc.add_paragraph(txt) 
        return
    # Initialize the recognizer 
    r = sr.Recognizer()      
    # Loop infinitely for user to 
    # speak 
    while(1): 
        with sr.Microphone() as source2 :
            # wait for a second to let the recognizer 
            # adjust the energy threshold based on 
            # the surrounding noise level 
            r.adjust_for_ambient_noise(source2, duration=0.2) 
            #listens for the user's input 
            audio2 = r.listen(source2)            
            # Using google to recognize audio 
            MyText = r.recognize_google(audio2) 
            MyText = MyText.lower()
            global s
            s.set(MyText)
            write(MyText) 
Label(myWindow, text="").pack()
Label(myWindow, text="").pack()
Label(myWindow, text="Enter Filename*").pack()
# Creating a photoimage object to use image 
photo = PhotoImage(file = r"C:\Users\Hp\Documents\Python Scripts\mic.png")
# Resizing image to fit on button 
photoimage = photo.subsample(3, 3)
name=StringVar()
e1 = Entry(myWindow,textvariable=name).pack()
Label(myWindow, text="").pack()
Label(myWindow, text="").pack()
# compound option is used to align image on LEFT side of button
mic=Button(myWindow, text = "Click to start Voice typing", image = photoimage,compound = LEFT,command=listen).pack()
Label(myWindow, text="").pack()
Label(myWindow, text="").pack()
s = StringVar()
e2 = Entry(myWindow,textvariable=s).pack()
Label(myWindow, text="").pack()
Label(myWindow, text="Enter the image location*").pack()
i = StringVar()
Entry(myWindow,textvariable=i).pack()
button2 = Button(text="Browse for image", command=browse_button).pack()
Label(myWindow, text="").pack()
Button(myWindow,text="insert image",command=insert).pack()
#creating the button for save option
Button(myWindow,text="save document",command=save).pack()
myWindow.mainloop()
